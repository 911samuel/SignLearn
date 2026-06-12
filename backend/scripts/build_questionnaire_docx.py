"""
Render the SignLearn agency questionnaire from Markdown into a clean .docx.

Output:
  artifacts/questionnaire/SignLearn_Agency_Questionnaire.docx
  ~/Downloads/SignLearn_Agency_Questionnaire.docx
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).parent.parent.parent
SRC = REPO / "artifacts" / "questionnaire" / "signlearn_agency_questionnaire.md"
OUT = REPO / "artifacts" / "questionnaire" / "SignLearn_Agency_Questionnaire.docx"
DOWNLOADS = Path.home() / "Downloads" / "SignLearn_Agency_Questionnaire.docx"


def add_runs(p, text: str):
    """Render **bold**, *italic*, `code` inline emphasis."""
    parts = re.split(r"(\*\*[^\*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            for sp in re.split(r"(\*[^\*]+\*|`[^`]+`)", part):
                if sp.startswith("*") and sp.endswith("*") and not sp.startswith("**"):
                    r = p.add_run(sp[1:-1]); r.italic = True
                elif sp.startswith("`") and sp.endswith("`"):
                    r = p.add_run(sp[1:-1]); r.font.name = "Courier New"
                else:
                    p.add_run(sp)


def main() -> None:
    doc = Document()

    # Page setup — A4, comfortable margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = SRC.read_text().splitlines()
    i = 0
    in_blockquote = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            # Horizontal rule → empty paragraph with a thin bottom border (simple visual break)
            p = doc.add_paragraph()
            p.add_run(" ").font.size = Pt(2)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(stripped[2:].strip())
            r.font.size = Pt(20); r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(stripped[3:].strip())
            r.font.size = Pt(14); r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x4F, 0xCB)
            i += 1; continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            r = p.add_run(stripped[4:].strip())
            r.font.size = Pt(12); r.bold = True
            i += 1; continue

        # Blockquote — italic indented paragraph
        if stripped.startswith(">"):
            body = stripped[1:].strip()
            # Strip leading bold markers
            if body.startswith("*") and body.endswith("*") and not body.startswith("**"):
                body = body[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            r = p.add_run(body); r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1; continue

        if not stripped:
            i += 1; continue

        # Bullet (☐ or -)
        if stripped.startswith("☐ ") or stripped.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            add_runs(p, stripped)
            i += 1; continue

        # Otherwise — paragraph (collapse soft-wrapped continuations)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", ">", "-", "☐", "|", "---"))
                    or nxt[:1].isdigit() and "." in nxt[:4]):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_runs(p, text)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    shutil.copy2(OUT, DOWNLOADS)
    print(f"Wrote {OUT}")
    print(f"Copy:  {DOWNLOADS}")
    print(f"Size:  {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
