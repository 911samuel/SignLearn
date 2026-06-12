"""
Inject the new and expanded thesis chapters into Final-Year Project.docx.

- Reads the existing docx (preserves all front-matter + Chapter 1 intact).
- Strips the old thin Chapter 2 body and old partial Chapter 3 body.
- Appends Chapter 2 (expanded), Chapter 3 (expanded), Chapters 4-5-6 (new).
- Inserts figures with captions from artifacts/thesis/figures/.
- Appends the 12 new references to the Bibliography.
- Saves as Final-Year Project.UPDATED.docx (does NOT overwrite original).

Run:
    python backend/scripts/inject_thesis_chapters.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).parent.parent.parent
DOWNLOADS = Path.home() / "Downloads"
SRC = DOWNLOADS / "Final-Year Project.docx"
DST = DOWNLOADS / "Final-Year Project.UPDATED.docx"

THESIS = REPO / "artifacts" / "thesis"
FIG_DIR = THESIS / "figures"

CH2_MD = THESIS / "chapter_2_literature.md"
CH3_MD = THESIS / "chapter_3_methodology.md"
CH4_MD = THESIS / "chapter_4_design.md"
CH5_MD = THESIS / "chapter_5_results.md"
CH6_MD = THESIS / "chapter_6_conclusion.md"
REFS_BIB = THESIS / "new_references.bib"

# Figure references inside markdown → actual files (filenames only, picked up from FIG_DIR)
FIGURE_MAP = {
    "fig_system_architecture.png": "fig_system_architecture.png",
    "fig_class_diagram.png": "fig_class_diagram.png",
    "fig_deployment.png": "fig_deployment.png",
    "fig_data_flow.png": "fig_data_flow.png",
    "fig_sequence_diagram.png": "fig_sequence_diagram.png",
    "fig_per_class_accuracy.png": "fig_per_class_accuracy.png",
    "fig_confusion_matrix.png": "fig_confusion_matrix.png",
    "fig_training_curves.png": "fig_training_curves.png",
    "fig_word_model_top5.png": "fig_word_model_top5.png",
    "fig_architecture_comparison.png": "fig_architecture_comparison.png",
    "fig_latency.png": "fig_latency.png",
}


def parse_markdown(text: str) -> list[tuple[str, str]]:
    """Convert markdown into a list of (kind, content) directives.

    Kinds:
      h1, h2, h3   – headings (Heading 1/2/3 style)
      para         – regular paragraph
      figure       – embed a figure by filename (parsed from '> *(see fig_X.png)*')
      caption      – figure caption ('> **Figure X.Y:** ...')
      table        – Markdown pipe table block
      bullet       – bullet list item
    """
    out: list[tuple[str, str]] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == "---":
            i += 1
            continue

        # Code fences (skip but keep contents as plain text)
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                out.append(("para", lines[i]))
                i += 1
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            out.append(("h3", stripped[4:].strip()))
            i += 1
            continue
        if stripped.startswith("## "):
            out.append(("h2", stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith("# "):
            out.append(("h1", stripped[2:].strip()))
            i += 1
            continue

        # Figure caption block: > **Figure ...** ... *(see fig_xxx.png)*
        if stripped.startswith("> **"):
            # Collect the whole blockquote (may span multiple lines).
            full = stripped[2:].strip()
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith(">"):
                full += " " + lines[j].strip()[1:].strip()
                j += 1
            # Extract figure reference if present
            m = re.search(r"\(see (fig_[\w\.]+)\)", full)
            fig_ref = m.group(1) if m else None
            # Strip the "(see fig_X.png)" trailer + surrounding italics from caption
            caption = re.sub(r"\s*\*?\(see fig_[\w\.]+\)\*?\s*", "", full).strip()
            if fig_ref:
                out.append(("figure", fig_ref))
            if caption:
                out.append(("caption", caption))
            i = j
            continue

        # Table block (starts with | …|, has separator row of dashes)
        if stripped.startswith("|") and "|" in stripped[1:]:
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            out.append(("table", "\n".join(block)))
            continue

        # Bullet
        if stripped.startswith("- "):
            out.append(("bullet", stripped[2:]))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Otherwise paragraph (collapse soft-wrapped lines into one para)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", ">", "-", "|")) or nxt == "---"
                    or nxt.startswith("```")):
                break
            para_lines.append(nxt)
            i += 1
        out.append(("para", " ".join(para_lines)))

    return out


def _add_runs_with_emphasis(p, text: str):
    """Split text on **bold** and *italic* and add appropriate runs."""
    # Tokenise on bold first
    parts = re.split(r"(\*\*[^\*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # Inside, also handle *italic* and `code`
            subparts = re.split(r"(\*[^\*]+\*|`[^`]+`)", part)
            for sp in subparts:
                if sp.startswith("*") and sp.endswith("*") and not sp.startswith("**"):
                    run = p.add_run(sp[1:-1])
                    run.italic = True
                elif sp.startswith("`") and sp.endswith("`"):
                    run = p.add_run(sp[1:-1])
                    run.font.name = "Courier New"
                else:
                    p.add_run(sp)


def add_paragraph(doc, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_runs_with_emphasis(p, text)
    return p


def add_figure(doc, fig_name: str):
    src = FIG_DIR / fig_name
    if not src.exists():
        print(f"  WARN: missing figure {src}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(src), width=Inches(6.0))


def add_caption(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs_with_emphasis(p, text)
    # Italicise the whole caption
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)


def add_table_from_markdown(doc, md_table: str):
    """Build a docx table from a markdown pipe-table block."""
    rows = [r.strip() for r in md_table.split("\n") if r.strip().startswith("|")]
    # Strip leading/trailing pipes
    def cells(row: str) -> list[str]:
        cs = row.strip().strip("|").split("|")
        return [c.strip() for c in cs]

    if len(rows) < 2:
        return
    header = cells(rows[0])
    # rows[1] is the separator line of dashes; skip it
    data_rows = [cells(r) for r in rows[2:] if not set(r.replace("|", "").replace(" ", "").replace(":", "").replace("-", "")) <= {""}]
    n_cols = len(header)

    tbl = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_runs_with_emphasis(p, h)
        for run in p.runs:
            run.bold = True
    for i, drow in enumerate(data_rows):
        for j in range(n_cols):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            txt = drow[j] if j < len(drow) else ""
            _add_runs_with_emphasis(cell.paragraphs[0], txt)
    # Spacing after table
    doc.add_paragraph()


def render_directives(doc, directives, *, start_at_h1: bool = True):
    """Render a parsed markdown stream into the document.

    If start_at_h1 is False, the top-level h1 (chapter title) is suppressed —
    used when the chapter heading is already present in the docx and we are
    appending body content.
    """
    skipped_first_h1 = False
    bullet_buffer = []
    def flush_bullets():
        nonlocal bullet_buffer
        for b in bullet_buffer:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_emphasis(p, b)
        bullet_buffer = []

    for kind, content in directives:
        if kind != "bullet":
            flush_bullets()
        if kind == "h1":
            if not start_at_h1 and not skipped_first_h1:
                skipped_first_h1 = True
                continue
            p = doc.add_paragraph(style="Heading 1")
            _add_runs_with_emphasis(p, content)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            _add_runs_with_emphasis(p, content)
        elif kind == "h3":
            p = doc.add_paragraph(style="Heading 3")
            _add_runs_with_emphasis(p, content)
        elif kind == "para":
            add_paragraph(doc, content)
        elif kind == "bullet":
            bullet_buffer.append(content)
        elif kind == "figure":
            add_figure(doc, FIGURE_MAP.get(content, content))
        elif kind == "caption":
            add_caption(doc, content)
        elif kind == "table":
            add_table_from_markdown(doc, content)
    flush_bullets()


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Source missing: {SRC}")

    doc = Document(str(SRC))

    # Strategy: we cannot reliably edit-in-place using python-docx without breaking
    # numbering. Safer is to APPEND new chapters AFTER the existing CHAPTER 3 body
    # and BEFORE the Bibliography. Then we *delete* the old thin Chapter 2 body
    # and the old thin Chapter 3 body to avoid duplication.
    #
    # Implementation: enumerate paragraphs, find indices of:
    #   - CHAPTER 2 heading
    #   - CHAPTER 3 heading
    #   - "Bibliography" heading
    # Then delete the bodies between Ch2 and Ch3, and between Ch3 and Bibliography.
    # Append new content before Bibliography.

    paragraphs = doc.paragraphs
    ch2_idx = ch3_idx = bib_idx = None
    for idx, p in enumerate(paragraphs):
        text = p.text.strip().upper()
        if text.startswith("CHAPTER 2") and ch2_idx is None:
            ch2_idx = idx
        elif text.startswith("CHAPTER 3") and ch3_idx is None:
            ch3_idx = idx
        elif text.startswith("BIBLIOGRAPHY") and bib_idx is None:
            bib_idx = idx
    if None in (ch2_idx, ch3_idx, bib_idx):
        raise SystemExit(
            f"Could not locate insertion anchors. ch2={ch2_idx} ch3={ch3_idx} bib={bib_idx}"
        )

    print(f"Anchors: Chapter 2 @ para {ch2_idx}, Chapter 3 @ para {ch3_idx}, "
          f"Bibliography @ para {bib_idx}")

    # Strategy: delete EVERYTHING from CH2 heading through (but not including)
    # Bibliography heading — both old chapter headings and their bodies. Then
    # stash Bibliography + trailing paragraphs. Append every chapter (Ch2..Ch6)
    # with its own H1. Restore Bibliography. Append new references.
    to_delete = list(paragraphs[ch2_idx:bib_idx])
    print(f"Deleting {len(to_delete)} paragraph(s) (old Ch2 heading+body, Ch3 heading+body)...")
    for p in to_delete:
        p._element.getparent().remove(p._element)

    # Stash Bibliography + everything after it
    paragraphs = doc.paragraphs
    bib_idx = next(i for i, p in enumerate(paragraphs)
                   if p.text.strip().upper().startswith("BIBLIOGRAPHY"))
    bib_para_xml = paragraphs[bib_idx]._element
    body = bib_para_xml.getparent()
    tail_elements = []
    p = bib_para_xml
    while p is not None:
        nxt = p.getnext()
        tail_elements.append(p)
        body.remove(p)
        p = nxt
    print(f"Stashed {len(tail_elements)} tail element(s) (bibliography + trailing).")

    # ---- Append all new chapters in order ----
    for ch_path, label in [
        (CH2_MD, "Chapter 2"),
        (CH3_MD, "Chapter 3"),
        (CH4_MD, "Chapter 4"),
        (CH5_MD, "Chapter 5"),
        (CH6_MD, "Chapter 6"),
    ]:
        print(f"Appending {label}...")
        _add_page_break(doc)
        dirs = parse_markdown(ch_path.read_text())
        render_directives(doc, dirs, start_at_h1=True)

    # ---- Restore Bibliography + tail, then append new refs ----
    print("Restoring Bibliography section...")
    _add_page_break(doc)
    body = doc.element.body
    for el in tail_elements:
        body.append(el)

    # Append the new references
    print("Appending 12 new IEEE references...")
    refs_text = REFS_BIB.read_text()
    for line in refs_text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        # Reference lines start with "[NN]" — others are continuations of the prior line.
        # We treat consecutive non-blank lines as one logical entry, separated by blank lines.
    # Simpler: split refs file on blank lines
    entries = []
    cur = []
    for line in refs_text.splitlines():
        if line.strip().startswith("#"):
            continue
        if line.strip() == "":
            if cur:
                entries.append(" ".join(s.strip() for s in cur))
                cur = []
        else:
            cur.append(line)
    if cur:
        entries.append(" ".join(s.strip() for s in cur))

    for entry in entries:
        if entry.startswith("["):
            add_paragraph(doc, entry)

    doc.save(str(DST))
    print(f"\nSaved: {DST}")
    print(f"Open in Word, press F9 to refresh fields (TOC, list of figures, table of contents).")


if __name__ == "__main__":
    main()
